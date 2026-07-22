"""Generate the editable DOCX title page (with author details) for the
Decision Support Systems submission."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
# base font
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)
    p.space_after = Pt(4)
    return p

# Title
t = doc.add_paragraph()
tr = t.add_run("SPADE: spline additive-noise DAG estimation for interpretable "
               "nonlinear causal discovery in managerial decision support")
tr.bold = True; tr.font.size = Pt(15)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Authors with superscript affiliation markers
authors = doc.add_paragraph(); authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
def author(name, sup, star=False):
    r = authors.add_run(name);
    s = authors.add_run(sup + ("," if not star else "")); s.font.superscript = True
    if star:
        a = authors.add_run("*"); a.font.superscript = True
    authors.add_run("  ")
authors.add_run("Quang-Vinh Dang")
r = authors.add_run("a,*"); r.font.superscript = True
authors.add_run(",  Dat Le")
r = authors.add_run("b"); r.font.superscript = True
authors.add_run(",  Minh Ngoc Dinh")
r = authors.add_run("c"); r.font.superscript = True

# Affiliations
aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
for mark, text in [("a", "British University Vietnam, Hung Yen, Vietnam"),
                   ("b", "Ho Chi Minh City University of Economics and Finance, Ho Chi Minh City, Vietnam"),
                   ("c", "Millennia Education, Ho Chi Minh City, Vietnam")]:
    s = aff.add_run(mark); s.font.superscript = True; s.italic = True
    rr = aff.add_run(" " + text + "\n"); rr.italic = True

# Corresponding author
corr = doc.add_paragraph()
corr.add_run("* Corresponding author. ").bold = True
corr.add_run("Email addresses: vinh.dq4@buv.edu.vn (Quang-Vinh Dang), "
             "datla@uef.edu.vn (Dat Le), minh.dinh@maeducation.com (Minh Ngoc Dinh).")

doc.add_paragraph()

# Abstract
heading("Abstract")
abstract = (
 "Understanding causal relationships in organizational time-series data is critical "
 "for managerial decision-making, yet forecasting systems often trade interpretability "
 "for accuracy. We introduce SPADE (SPline Additive-noise DAG Estimation), an "
 "information-systems artifact built on Kolmogorov–Arnold networks that jointly "
 "discovers interpretable non-linear causal structure and produces one-step forecasts "
 "from a single differentiable model. Its headline result is on non-linear instantaneous "
 "causal discovery, where learnable spline edges recover additive-noise DAGs at "
 "AUROC ≈ 0.88 across d ∈ {6, 10, 20}, decisively outperforming the strongest "
 "modern DAG learners (DAGMA-nonlinear, GraN-DAG, NOTEARS-MLP). We ground the design in "
 "the identifiability theory of additive-noise models, evaluate transparently against "
 "recent baselines with threshold-free metrics, and report all limitations candidly.")
doc.add_paragraph(abstract)

# Keywords
kw = doc.add_paragraph()
kw.add_run("Keywords: ").bold = True
kw.add_run("Causal discovery; Kolmogorov–Arnold networks; Interpretable machine "
           "learning; Time-series forecasting; Decision support systems; Directed acyclic graphs")

doc.add_paragraph()

# Declarations
heading("Declarations")
d1 = doc.add_paragraph(); d1.add_run("Funding. ").bold = True
d1.add_run("This research received no specific grant from funding agencies in the public, "
           "commercial, or not-for-profit sectors. [Authors: amend if any funding applies.]")
d2 = doc.add_paragraph(); d2.add_run("Conflict of interest. ").bold = True
d2.add_run("The authors declare no competing interests.")
d3 = doc.add_paragraph(); d3.add_run("Data and code availability. ").bold = True
d3.add_run("Code and data to reproduce all results are available in the accompanying repository.")

doc.save("title_page.docx")
print("wrote title_page.docx")
